"""
ClauseGuard - AI-Powered Contract Risk Analysis
Production-ready backend using Flask + Groq API
"""

import os
import json
import re
import io
import uuid
import hashlib
import tempfile
from datetime import datetime
from flask import Flask, request, jsonify, render_template, send_file, session
from openai import OpenAI
from prompts import SYSTEM_PROMPT, ANALYSIS_PROMPT, QUICK_CLAUSE_PROMPT, EDIT_SUGGESTION_PROMPT

app = Flask(__name__)
app.secret_key = os.environ.get("SECRET_KEY", "clauseguard-dev-key-change-in-prod")
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max upload

# Groq API setup (OpenAI-compatible, free tier)
GROQ_API_KEY = os.environ.get("GROQ_API_KEY", "")
GROQ_BASE_URL = "https://api.groq.com/openai/v1"
GROQ_MODEL = os.environ.get("GROQ_MODEL", "llama-3.3-70b-versatile")

client = None
if GROQ_API_KEY:
    client = OpenAI(api_key=GROQ_API_KEY, base_url=GROQ_BASE_URL)

# In-memory storage for contract history (use database in production)
contract_history = {}

# In-memory storage for original uploaded files (keyed by session ID)
uploaded_files = {}

LANG_INSTRUCTION = {
    'zh': '\n\nIMPORTANT: You MUST respond with all text fields (risk_explanation, suggested_revision, revision_rationale, executive_summary, negotiation_tips, legal_notes, risk_change, enforceability_notes, legal_reference) in Chinese (Simplified). Keep JSON keys, risk levels (HIGH/MEDIUM/LOW), and party names in English. The original_text should remain as-is from the contract.',
    'en': ''
}


def call_ai(system_prompt, user_prompt, temperature=0.1):
    """Call Groq API with expert contract prompts."""
    if not client:
        return {"error": "GROQ_API_KEY not set. Get one free at https://console.groq.com/keys"}

    try:
        response = client.chat.completions.create(
            model=GROQ_MODEL,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            temperature=temperature,
            max_tokens=8000,
            response_format={"type": "json_object"}
        )
        content = response.choices[0].message.content
        return json.loads(content)
    except json.JSONDecodeError:
        try:
            match = re.search(r'\{.*\}', content, re.DOTALL)
            if match:
                return json.loads(match.group())
        except Exception:
            pass
        return {"error": "Failed to parse AI response", "raw": content}
    except Exception as e:
        return {"error": str(e)}


def extract_text_from_bytes(file_bytes, filename):
    """Extract text from file bytes (PDF, DOCX, or TXT)."""
    filename = filename.lower()

    if filename.endswith('.txt'):
        return file_bytes.decode('utf-8')

    elif filename.endswith('.pdf'):
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(io.BytesIO(file_bytes))
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            return text
        except Exception as e:
            return f"Error reading PDF: {str(e)}"

    elif filename.endswith('.docx'):
        try:
            from docx import Document
            doc = Document(io.BytesIO(file_bytes))
            return "\n".join([p.text for p in doc.paragraphs])
        except Exception as e:
            return f"Error reading DOCX: {str(e)}"

    return "Unsupported file format. Please upload PDF, DOCX, or TXT."


def get_session_id():
    """Get or create a session ID for tracking contract history."""
    if 'session_id' not in session:
        session['session_id'] = str(uuid.uuid4())
    return session['session_id']


# ============ ROUTES ============

@app.route('/')
def index():
    has_key = bool(GROQ_API_KEY)
    return render_template('index.html', has_api_key=has_key)


@app.route('/api/analyze', methods=['POST'])
def analyze_contract():
    """Full contract risk analysis."""
    contract_text = ""
    lang = 'en'
    sid = get_session_id()

    if 'file' in request.files and request.files['file'].filename:
        uploaded_file = request.files['file']
        orig_filename = uploaded_file.filename
        filename_lower = orig_filename.lower()
        file_bytes = uploaded_file.read()

        # Store original file for later export
        uploaded_files[sid] = {
            'bytes': file_bytes,
            'filename': orig_filename,
            'type': 'pdf' if filename_lower.endswith('.pdf') else 'docx' if filename_lower.endswith('.docx') else 'txt'
        }

        # Extract text from the stored bytes
        contract_text = extract_text_from_bytes(file_bytes, orig_filename)
        lang = request.form.get('lang', 'en')
    elif request.json and 'text' in request.json:
        contract_text = request.json['text']
        lang = request.json.get('lang', 'en')
        # Store as txt for pasted text
        uploaded_files[sid] = {
            'bytes': contract_text.encode('utf-8'),
            'filename': 'contract.txt',
            'type': 'txt'
        }
    else:
        return jsonify({"error": "No contract text provided"}), 400

    if len(contract_text.strip()) < 50:
        return jsonify({"error": "Contract text too short. Please provide a complete contract."}), 400

    prompt = ANALYSIS_PROMPT.replace("{contract_text}", contract_text)
    lang_suffix = LANG_INSTRUCTION.get(lang, '')
    result = call_ai(SYSTEM_PROMPT + lang_suffix, prompt)

    if "error" in result:
        return jsonify(result), 500

    # Save to history
    sid = get_session_id()
    if sid not in contract_history:
        contract_history[sid] = []

    history_entry = {
        "id": str(uuid.uuid4())[:8],
        "timestamp": datetime.utcnow().isoformat(),
        "contract_type": result.get("contract_type", "Unknown"),
        "risk_score": result.get("overall_risk_score", 0),
        "parties": result.get("parties", []),
        "clause_count": len(result.get("clauses", [])),
        "high_risks": len([c for c in result.get("clauses", []) if c.get("risk_level") == "HIGH"]),
        "lang": lang
    }
    contract_history[sid].append(history_entry)

    # Add history ID to result
    result["history_id"] = history_entry["id"]

    return jsonify(result)


@app.route('/api/analyze-clause', methods=['POST'])
def analyze_clause():
    """Analyze a single selected clause."""
    data = request.json
    clause_text = data.get('clause_text', '')
    contract_type = data.get('contract_type', 'General Agreement')
    lang = data.get('lang', 'en')

    if not clause_text:
        return jsonify({"error": "No clause text provided"}), 400

    prompt = QUICK_CLAUSE_PROMPT.replace("{clause_text}", clause_text).replace("{contract_type}", contract_type)
    lang_suffix = LANG_INSTRUCTION.get(lang, '')
    result = call_ai(SYSTEM_PROMPT + lang_suffix, prompt)
    return jsonify(result)


@app.route('/api/suggest-edit', methods=['POST'])
def suggest_edit():
    """Get AI-powered edit suggestion for a clause."""
    data = request.json
    original_text = data.get('original_text', '')
    user_intent = data.get('user_intent', 'Make this clause more balanced and protective')
    contract_type = data.get('contract_type', 'General Agreement')
    lang = data.get('lang', 'en')

    if not original_text:
        return jsonify({"error": "No clause text provided"}), 400

    prompt = (EDIT_SUGGESTION_PROMPT
              .replace("{original_text}", original_text)
              .replace("{user_intent}", user_intent)
              .replace("{contract_type}", contract_type))
    lang_suffix = LANG_INSTRUCTION.get(lang, '')
    result = call_ai(SYSTEM_PROMPT + lang_suffix, prompt)
    return jsonify(result)


@app.route('/api/compare', methods=['POST'])
def compare_clauses():
    """Compare original vs revised clause text."""
    data = request.json
    original = data.get('original', '')
    revised = data.get('revised', '')

    if not original or not revised:
        return jsonify({"error": "Both original and revised text required"}), 400

    # Simple word-level diff
    orig_words = original.split()
    rev_words = revised.split()

    diff_html = []
    import difflib
    matcher = difflib.SequenceMatcher(None, orig_words, rev_words)
    for op, i1, i2, j1, j2 in matcher.get_opcodes():
        if op == 'equal':
            diff_html.append(' '.join(orig_words[i1:i2]))
        elif op == 'delete':
            diff_html.append(f'<del>{" ".join(orig_words[i1:i2])}</del>')
        elif op == 'insert':
            diff_html.append(f'<ins>{" ".join(rev_words[j1:j2])}</ins>')
        elif op == 'replace':
            diff_html.append(f'<del>{" ".join(orig_words[i1:i2])}</del>')
            diff_html.append(f'<ins>{" ".join(rev_words[j1:j2])}</ins>')

    return jsonify({"diff_html": ' '.join(diff_html)})


@app.route('/api/history', methods=['GET'])
def get_history():
    """Get contract analysis history for current session."""
    sid = get_session_id()
    history = contract_history.get(sid, [])
    return jsonify({"history": list(reversed(history))})


@app.route('/api/export', methods=['POST'])
def export_contract():
    """Export the contract with edits applied to the ORIGINAL uploaded file."""
    data = request.json
    edits = data.get('edits', [])  # [{old_text, new_text}, ...]

    sid = get_session_id()
    original = uploaded_files.get(sid)

    if not original:
        return jsonify({"error": "No original file found. Please re-upload and analyze first."}), 400

    file_type = original['type']
    file_bytes = original['bytes']
    filename = original['filename']

    # If no edits, return the original file as-is
    if not edits:
        buf = io.BytesIO(file_bytes)
        if file_type == 'pdf':
            return send_file(buf, mimetype='application/pdf', as_attachment=True,
                             download_name=filename)
        elif file_type == 'docx':
            return send_file(buf, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                             as_attachment=True, download_name=filename)
        else:
            return send_file(buf, mimetype='text/plain', as_attachment=True, download_name=filename)

    try:
        if file_type == 'docx':
            return export_modified_docx(file_bytes, edits, filename)
        elif file_type == 'pdf':
            return export_modified_pdf(file_bytes, edits, filename)
        else:
            return export_modified_txt(file_bytes, edits, filename)
    except Exception as e:
        return jsonify({"error": f"Export failed: {str(e)}"}), 500


def normalize_whitespace(text):
    """Normalize whitespace for fuzzy matching."""
    return re.sub(r'\s+', ' ', text.strip())


def has_cjk(text):
    """Check if text contains Chinese/Japanese/Korean characters."""
    return bool(re.search(r'[\u4e00-\u9fff\u3040-\u309f\u30a0-\u30ff]', text))


def find_best_match_paragraphs(paragraphs, old_text):
    """Find paragraphs that contain the old text, handling multi-paragraph clauses.
    Returns list of (index, paragraph) tuples."""
    old_normalized = normalize_whitespace(old_text)
    matches = []

    # First try: exact single paragraph match
    for i, para in enumerate(paragraphs):
        para_text = normalize_whitespace(para.text)
        if para_text and para_text == old_normalized:
            return [(i, para)]

    # Second try: paragraph text is a substantial substring of old_text
    for i, para in enumerate(paragraphs):
        para_text = normalize_whitespace(para.text)
        if para_text and len(para_text) > 20 and para_text in old_normalized:
            matches.append((i, para))

    # Third try: old_text is substring of paragraph
    if not matches:
        for i, para in enumerate(paragraphs):
            para_text = normalize_whitespace(para.text)
            if para_text and len(para_text) > 20 and old_normalized in para_text:
                matches.append((i, para))

    # Fourth try: first 50 chars match
    if not matches:
        old_start = old_normalized[:50]
        for i, para in enumerate(paragraphs):
            para_text = normalize_whitespace(para.text)
            if para_text and old_start in para_text:
                matches.append((i, para))

    return matches


def replace_paragraph_text(paragraph, new_text):
    """Replace paragraph text while preserving the formatting of existing runs."""
    if not paragraph.runs:
        paragraph.text = new_text
        return

    # Preserve the formatting (font, size, bold, italic, color) from the first run
    first_run = paragraph.runs[0]
    first_run.text = new_text

    # Clear all other runs (keep XML elements but empty text)
    for run in paragraph.runs[1:]:
        run.text = ""


def export_modified_docx(file_bytes, edits, filename):
    """Modify the original DOCX in-place, preserving formatting."""
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))

    # Collect all paragraphs including those in tables
    all_paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_paragraphs.extend(cell.paragraphs)

    for edit in edits:
        old_text = edit.get('old_text', '')
        new_text = edit.get('new_text', '')
        if not old_text or not new_text or old_text == new_text:
            continue

        matches = find_best_match_paragraphs(all_paragraphs, old_text)

        if len(matches) == 1:
            # Single paragraph match — replace its text
            _, para = matches[0]
            replace_paragraph_text(para, new_text)
        elif len(matches) > 1:
            # Multiple paragraphs matched (clause spans multiple paragraphs)
            # Put all new text in the first matched paragraph, clear the rest
            _, first_para = matches[0]
            replace_paragraph_text(first_para, new_text)
            for _, para in matches[1:]:
                replace_paragraph_text(para, "")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    return send_file(buf,
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                     as_attachment=True,
                     download_name=filename)


def export_modified_pdf(file_bytes, edits, filename):
    """Modify the original PDF in-place using PyMuPDF, preserving layout."""
    import fitz  # PyMuPDF

    doc = fitz.open(stream=file_bytes, filetype="pdf")

    for edit in edits:
        old_text = edit.get('old_text', '')
        new_text = edit.get('new_text', '')
        if not old_text or not new_text or old_text == new_text:
            continue

        # Try multiple search snippets for robustness
        search_attempts = [
            old_text[:80].strip(),
            ' '.join(old_text.split()[:10]),
            ' '.join(old_text.split()[:5]),
        ]

        for page in doc:
            text_instances = None
            for snippet in search_attempts:
                text_instances = page.search_for(snippet)
                if text_instances:
                    break

            if not text_instances:
                continue

            # Get font info from the matched area
            first_rect = text_instances[0]
            blocks = page.get_text("dict", clip=first_rect).get("blocks", [])
            font_size = 11
            font_name = "helv"
            text_color = (0, 0, 0)

            for block in blocks:
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        font_size = span.get("size", 11)
                        color_int = span.get("color", 0)
                        r = ((color_int >> 16) & 0xFF) / 255
                        g = ((color_int >> 8) & 0xFF) / 255
                        b = (color_int & 0xFF) / 255
                        text_color = (r, g, b)
                        break
                    break
                break

            # Build the combined area of all matched text rects
            combined = text_instances[0]
            for r in text_instances[1:]:
                combined = combined | r

            # Use page margins from the first rect position
            page_rect = page.rect
            combined.x0 = max(page_rect.x0 + 20, combined.x0 - 10)
            combined.x1 = min(page_rect.x1 - 20, combined.x1 + 10)

            # Redact (white out) the old text
            page.add_redact_annot(combined, fill=(1, 1, 1))
            page.apply_redactions()

            # Choose font: use CJK font for Chinese text
            if has_cjk(new_text):
                font_name = "china-ss"  # PyMuPDF built-in Simplified Chinese font

            # Insert new text at the same position with extra height for reflow
            extra_height = max(20, (len(new_text) // 40) * font_size)
            text_rect = fitz.Rect(combined.x0, combined.y0,
                                  combined.x1, combined.y1 + extra_height)
            page.insert_textbox(text_rect, new_text,
                                fontsize=font_size,
                                fontname=font_name,
                                color=text_color,
                                align=fitz.TEXT_ALIGN_LEFT)
            break  # Only process first page where text is found

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    doc.close()

    return send_file(buf, mimetype='application/pdf', as_attachment=True,
                     download_name=filename)


def export_modified_txt(file_bytes, edits, filename):
    """Modify plain text file with find-and-replace."""
    text = file_bytes.decode('utf-8')

    for edit in edits:
        old_text = edit.get('old_text', '')
        new_text = edit.get('new_text', '')
        if old_text and new_text and old_text != new_text:
            # Try exact match first, then normalized match
            if old_text in text:
                text = text.replace(old_text, new_text)
            else:
                # Fuzzy: normalize whitespace and try again
                old_norm = normalize_whitespace(old_text)
                text_norm = normalize_whitespace(text)
                if old_norm in text_norm:
                    text = text.replace(old_norm, new_text)

    buf = io.BytesIO(text.encode('utf-8'))
    return send_file(buf, mimetype='text/plain', as_attachment=True, download_name=filename)


@app.route('/api/file-type', methods=['GET'])
def get_file_type():
    """Get the type of the original uploaded file."""
    sid = get_session_id()
    original = uploaded_files.get(sid)
    if not original:
        return jsonify({"type": "none"})
    return jsonify({"type": original['type'], "filename": original['filename']})


@app.route('/health')
def health():
    """Health check endpoint for deployment."""
    return jsonify({
        "status": "healthy",
        "model": GROQ_MODEL,
        "api_connected": bool(GROQ_API_KEY)
    })


if __name__ == '__main__':
    if not GROQ_API_KEY:
        print("\n" + "=" * 60)
        print("  WARNING: GROQ_API_KEY not set!")
        print("  Get your free key at: https://console.groq.com/keys")
        print("  Then run: GROQ_API_KEY=your_key python3 app.py")
        print("=" * 60 + "\n")
    else:
        print("\n  ClauseGuard AI is ready!")
        print(f"  Using model: {GROQ_MODEL}\n")

    port = int(os.environ.get("PORT", 5000))
    app.run(debug=True, port=port)
